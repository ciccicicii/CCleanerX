from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "course_deliverables"
OUT.mkdir(exist_ok=True)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(102, 112, 133)
LIGHT = "F2F4F7"
PALE = "E8EEF5"
BORDER = "D0D5DD"
BODY_FONT = "Microsoft YaHei UI"
MONO_FONT = "Consolas"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_run_font(run, name=BODY_FONT, size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_p(doc, text="", bold=False, color=None, size=None, align=None, after=6, before=0, font=BODY_FONT):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    if align is not None:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, font, size=size or 11, color=color or INK, bold=bold)
    return paragraph


def add_bullet(doc, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, BODY_FONT, size=10.5, color=INK)


def add_number(doc, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, BODY_FONT, size=10.5, color=INK)


def add_code(doc, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, MONO_FONT, size=9.2, color=RGBColor(52, 64, 84))


def add_heading(doc, text: str, level=1):
    paragraph = doc.add_heading("", level=level)
    paragraph.paragraph_format.space_before = Pt({1: 16, 2: 12, 3: 8}.get(level, 6))
    paragraph.paragraph_format.space_after = Pt({1: 8, 2: 6, 3: 4}.get(level, 4))
    run = paragraph.add_run(text)
    set_run_font(
        run,
        BODY_FONT,
        size={1: 16, 2: 13, 3: 12}.get(level, 11),
        color=BLUE if level < 3 else DARK_BLUE,
        bold=True,
    )


def add_table(doc, headers: list[str], rows: list[tuple[str, ...]], widths=None, header_fill=LIGHT):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_shading(header_cells[index], header_fill)
        paragraph = header_cells[index].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, BODY_FONT, size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, BODY_FONT, size=9.2, color=INK)
    if widths is None:
        base = 9360 // len(headers)
        widths = [base] * len(headers)
        widths[-1] += 9360 - sum(widths)
    set_table_width(table, widths)
    add_p(doc, "", after=4)
    return table


def add_callout(doc, title: str, body: str, fill="F4F6F9") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_run_font(run, BODY_FONT, size=10.5, color=DARK_BLUE, bold=True)
    body_paragraph = cell.add_paragraph()
    body_paragraph.paragraph_format.space_after = Pt(0)
    body_run = body_paragraph.add_run(body)
    set_run_font(body_run, BODY_FONT, size=10, color=INK)
    set_table_width(table, [9360])
    add_p(doc, "", after=4)


def setup_doc(title: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    header = section.header.paragraphs[0]
    header.text = title
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.runs[0], BODY_FONT, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "本地磁盘清理器课程设计成果文档"
    set_run_font(footer.runs[0], BODY_FONT, size=9, color=MUTED)
    return doc


def add_cover(doc, title: str, subtitle: str) -> None:
    add_p(doc, "互联网应用开发实践", bold=True, color=MUTED, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=16, before=24)
    add_p(doc, title, bold=True, color=INK, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_p(doc, subtitle, color=MUTED, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    rows = [
        ("系（院）", "（请填写）"),
        ("专业班级", "网络工程（请填写）"),
        ("姓名", "（请填写）"),
        ("学号", "（请填写）"),
        ("指导教师", "向华"),
        ("项目名称", "本地磁盘清理器 Local Disk Cleaner"),
        ("设计时间", "2026年6月"),
    ]
    table = add_table(doc, ["项目", "内容"], rows, widths=[2200, 7160], header_fill=PALE)
    for row in table.rows[1:]:
        set_cell_shading(row.cells[0], "F9FAFB")
        row.cells[0].paragraphs[0].runs[0].bold = True
    add_callout(doc, "项目定位", "本项目面向 Windows 本地磁盘空间不足场景，提供磁盘扫描、分类分析、安全清理、AI 解释建议和可打包分发的本地桌面应用。")
    doc.add_page_break()


def add_static_toc(doc, items: list[str]) -> None:
    add_heading(doc, "目录", 1)
    for item in items:
        add_p(doc, item, after=2)
    doc.add_page_break()


def build_report() -> None:
    doc = setup_doc("课程设计报告")
    add_cover(doc, "本地磁盘清理器课程设计报告", "Local Disk Cleaner - 安全边界驱动的本地磁盘分析与清理工具")
    add_static_toc(doc, ["一、课程设计目的", "二、设计任务及要求", "三、需求分析", "四、总体设计", "五、详细设计与实现", "六、系统测试与结果分析", "七、AI工具与Git工程化实践", "八、课程设计小结"])

    add_heading(doc, "一、课程设计目的", 1)
    add_p(doc, "本课程设计围绕“互联网应用开发实践”的综合训练目标，完成一款面向个人电脑磁盘空间管理场景的本地应用软件。项目重点训练需求分析、系统架构设计、数据处理、可视化展示、安全控制、测试验证和工程化交付能力。")
    add_p(doc, "本项目选择“本地磁盘清理器”作为课设题目，是因为磁盘空间不足是普通用户和开发者常见问题。传统清理工具往往只给出笼统删除建议，用户难以判断哪些文件可以删、哪些文件应谨慎处理。本项目通过规则驱动分类、回收站安全清理和 AI 解释层，将“找出占用空间大户”和“给出可执行处置建议”结合起来。")
    add_callout(doc, "培养目标对应", "项目对应毕业要求 3.3 中的整体方案设计、系统安全、运行健康和合规伦理考量；也对应毕业要求 11.2 中的前沿工具跟踪、问题分析、迭代优化与总结能力。")

    add_heading(doc, "二、设计任务及要求", 1)
    add_p(doc, "根据任务书要求，本项目需独立完成互联网应用软件开发实践，提交完整课程设计说明书、源代码、配置文件、运行说明、Git 版本记录、运行演示和测试记录。结合项目实际，任务拆解如下：")
    add_table(doc, ["任务要求", "本项目对应实现"], [
        ("项目选题与需求分析", "围绕 Windows 本地磁盘空间不足场景，明确目标用户、功能模块和安全边界。"),
        ("系统方案创新设计", "设计扫描器、规则分析器、清理安全层、桌面 UI、AI 深度分析解释层。"),
        ("专项开发任务", "完成数据处理与可视化、轻量化运维工具、智能功能拓展三类专项任务。"),
        ("工程化规范", "使用 Python 模块化开发、单元测试、PyInstaller 打包、Git 版本管理。"),
        ("测试与优化", "围绕扫描、分类、清理权限、防伪 action_id、AI JSON 容错、桌面 UI 等建立自动化测试。"),
    ], widths=[2500, 6860])

    add_heading(doc, "三、需求分析", 1)
    add_heading(doc, "3.1 用户与场景", 2)
    add_p(doc, "目标用户包括 Windows 普通用户、开发者和需要临时释放磁盘空间的学生。典型场景包括 C 盘空间不足、游戏库占用过大、下载目录堆积、AppData 缓存增长、Python/Node/浏览器缓存占用异常等。")
    add_heading(doc, "3.2 功能需求", 2)
    for item in ["选择磁盘并发起扫描，扫描过程保持只读。", "统计大文件和目录，按占用大小排序。", "将结果划分为可安全清理、可清理但需确认、需要人工判断、谨慎处理、其他占用。", "对绿色项目提供“移入回收站”清理动作，对其他项目主要提供打开位置和说明。", "支持 AI 智能深度分析，进一步扫描人工判断目录的子项，并生成解释建议。", "支持打包为 Windows 本地 exe，双击启动桌面应用。"]:
        add_bullet(doc, item)
    add_heading(doc, "3.3 非功能需求与约束", 2)
    for item in ["安全性：UI 不能提交任意路径，清理动作只能由本地规则生成的 action_id 触发。", "可靠性：遇到权限拒绝、文件占用、WinError 32 等情况时跳过并记录，不中断整体流程。", "隐私性：发送给模型的数据尽量使用路径别名，不让模型决定删除权限。", "易用性：界面采用盘符选择、分类卡片、项目列表和分析说明的工作流。", "可维护性：扫描、分析、清理、AI、桌面 UI、报告渲染分模块实现。"]:
        add_bullet(doc, item)

    add_heading(doc, "四、总体设计", 1)
    add_heading(doc, "4.1 系统架构", 2)
    add_code(doc, "LocalDiskCleaner.exe\n  -> desktop_app.py 桌面窗口\n     -> scanner.py 只读扫描磁盘占用\n     -> analyzer.py 本地规则分类与 action_id 生成\n     -> ai_deep_scan.py 深度扫描人工判断目录\n     -> model_explainer.py 调用 DeepSeek 生成解释建议\n     -> cleanup.py 校验 action_id 并移入回收站\n     -> report.py / webapp.py 生成 HTML 报告或本地服务")
    add_heading(doc, "4.2 数据流设计", 2)
    add_code(doc, "scan() -> ScanResult\nanalyze(ScanResult) -> AnalysisResult\nrender(AnalysisResult) -> report.html / 桌面列表\nAI deep scan -> 合并新的候选项目 -> Model Explanation\ncleanup(action_id) -> resolve_action() -> validate -> recycle_bin/open")
    add_heading(doc, "4.3 模块划分", 2)
    add_table(doc, ["模块", "主要文件", "职责"], [
        ("扫描模块", "tools/scanner.py", "枚举磁盘和热点目录，统计大小、文件数、权限拒绝和占用情况。"),
        ("规则分析模块", "tools/analyzer.py", "按本地规则分类，生成绿色、黄色、红色、其他项目。"),
        ("清理安全模块", "tools/cleanup.py", "解析 action_id，验证层级和路径边界，执行打开位置或移入回收站。"),
        ("桌面应用模块", "tools/desktop_app.py", "提供原生 Tkinter 窗口、盘符选择、扫描按钮、分类列表、AI 说明和操作按钮。"),
        ("AI 分析模块", "tools/ai_deep_scan.py, tools/model_explainer.py", "对人工判断目录进行深度扫描，并请求模型生成解释层建议。"),
        ("报告与服务模块", "tools/report.py, tools/server.py, tools/webapp.py", "生成 HTML 报告，支持本地服务和浏览器交互版本。"),
        ("打包模块", "packaging/build_exe.ps1", "使用 PyInstaller 生成可分发 exe。"),
    ], widths=[1700, 2600, 5060])

    add_heading(doc, "五、详细设计与实现", 1)
    add_heading(doc, "5.1 只读扫描实现", 2)
    add_p(doc, "扫描模块使用 os.scandir 递归统计目录大小，并记录 denied_count、locked_count、skipped_reparse_count、file_count、dir_count。为避免误入符号链接或 Windows junction，扫描前会跳过 reparse point。遇到 WinError 32 文件占用时只计数并继续扫描。")
    add_code(doc, "def size_tree(path: str) -> SizeStats:\n    if os.path.islink(path) or is_reparse_point(path):\n        stats.skipped_reparse_count += 1\n        return stats\n    ...\n    except OSError as exc:\n        if is_windows_locked(exc):\n            stats.locked_count += 1")
    add_heading(doc, "5.2 规则分类实现", 2)
    add_p(doc, "分析模块将目录和文件按规则映射为不同风险等级。绿色规则只覆盖可重新生成的缓存和临时文件；黄色规则覆盖桌面、下载、文档、聊天软件数据等需要人工判断的内容；红色规则覆盖 Windows、Program Files、磁盘根目录等不应手动删除的位置。")
    add_table(doc, ["分类", "含义", "典型路径", "允许动作"], [
        ("可安全清理", "已通过本地安全校验，可移入回收站", "%TEMP%、浏览器缓存、NVIDIA DXCache、pip/npm 缓存", "打开位置、移入回收站"),
        ("可清理但需确认", "AI 深扫发现可能可清理，但需用户确认", "安装包、游戏安装目录、旧下载文件", "打开位置"),
        ("需要人工判断", "可能包含用户文件或应用数据", "Desktop、Documents、Downloads、Tencent/WeChat 数据", "打开位置"),
        ("谨慎处理", "系统或程序目录，不建议手动删除", "Windows、Program Files、ProgramData", "说明或打开位置"),
        ("其他占用", "已统计但无明确建议", "未知目录或普通数据目录", "查看说明"),
    ], widths=[1600, 2300, 3560, 1900])
    add_heading(doc, "5.3 清理安全边界", 2)
    add_p(doc, "清理层采用 action_id 机制。桌面 UI 或网页报告不会把路径直接传给清理模块，而是提交本次分析结果中已经生成的 action_id。cleanup.py 根据 action_id 找到项目，再校验 tier、kind、路径边界和禁止删除目录。只有 green 项的 recycle_bin 动作可以执行真实清理。")
    add_callout(doc, "安全原则", "模型不能授予删除权限，UI 不能提交任意路径，红色和黄色项目不能伪造回收站 action。所有真实清理默认移入回收站，不做永久删除。")
    add_heading(doc, "5.4 桌面界面实现", 2)
    add_p(doc, "桌面端采用 Tkinter 实现，打包后为 Windows 原生 exe。界面顶部为盘符选择与操作按钮，中部为分类概览卡片和项目列表，右侧为 AI 分析说明、项目详情和操作按钮。用户流程为：选择盘符 -> 开始扫描 -> 查看分类 -> 智能深度分析 -> 打开位置或移入回收站。")
    add_heading(doc, "5.5 AI 智能深度分析", 2)
    add_p(doc, "第二版加入 DeepSeek 作为解释层。AI 深度分析首先由本地代码继续扫描黄色/人工判断目录中的子项，识别可本地验证的缓存和需确认候选，然后将结果合并回主列表。DeepSeek 只负责解释结果、提示风险和给出优先检查建议，不能生成 action_id、allowlist 或删除权限。")

    add_heading(doc, "六、系统测试与结果分析", 1)
    add_p(doc, "2026-06-16 最近一次全量测试：python -m unittest discover -s tests -v，运行 63 项测试，结果 OK。")
    add_table(doc, ["测试类别", "覆盖内容", "结果"], [
        ("扫描测试", "磁盘根目录规范化、热点目录选择、扫描大小统计、权限异常处理", "通过"),
        ("规则分类测试", "绿色缓存、黄色用户数据、红色系统路径、未知路径分类", "通过"),
        ("清理安全测试", "伪造 action_id、黄色/红色回收站动作拒绝、缺失路径、文件占用部分成功", "通过"),
        ("AI 分析测试", "深扫候选收集、绿色候选合并、模型返回 JSON 容错、解释权限隔离", "通过"),
        ("桌面 UI 测试", "盘符选择、分类卡片、按钮文案、详情格式和分析说明格式", "通过"),
        ("打包验证", "PyInstaller 输出 LocalDiskCleaner.exe，zip 包包含主程序", "通过"),
    ], widths=[1900, 5660, 1800])

    add_heading(doc, "七、AI工具与Git工程化实践", 1)
    add_p(doc, "开发过程中使用 Codex 辅助完成需求拆解、测试编写、代码实现、异常排查、UI 优化、文档整理和打包验证。所有 AI 输出均经过本地运行、单元测试和人工确认后才整合进项目。")
    add_p(doc, "本地工作区存在 Git 仓库，但当前环境未配置远程仓库地址，git log 未返回可展示的提交记录。因此提交材料中应补充实际提交截图或在上传 GitHub/Gitee 后填写远程链接。本文档的“Git版本管理与迭代记录”已整理出可用于提交的阶段性迭代说明。")

    add_heading(doc, "八、课程设计小结", 1)
    add_p(doc, "本项目从最初的脚本扫描和 HTML 报告，逐步迭代为可选择盘符、可智能深度分析、可安全回收清理、可打包分发的本地桌面应用。开发过程中最大的难点在于删除权限安全边界：如果允许 UI 或 AI 直接决定删除路径，虽然功能更“自动”，但风险很高。因此最终采用规则层生成 action_id、清理层二次校验、模型只做解释的架构。")
    add_p(doc, "通过本课程设计，我进一步理解了互联网应用开发不仅是界面和功能实现，还包括需求边界、用户风险、数据隐私、错误处理、工程化测试和可交付包装。后续可继续优化方向包括：引入更完善的日志系统、加入更细粒度的磁盘可视化图表、支持应用卸载器跳转、提供多语言界面，以及在公开发布前改用后端代理或用户自配 API Key。")
    doc.save(OUT / "01-local-disk-cleaner-course-report.docx")


def build_run_guide() -> None:
    doc = setup_doc("项目运行说明")
    add_cover(doc, "项目运行说明", "本地磁盘清理器 Local Disk Cleaner")
    add_heading(doc, "一、项目文件结构", 1)
    add_table(doc, ["路径", "说明"], [
        ("tools/", "核心程序模块，包括扫描、分析、清理、AI、桌面 UI 和报告服务。"),
        ("tests/", "单元测试与测试夹具。"),
        ("packaging/", "PyInstaller 打包脚本和 spec 文件。"),
        ("dist/LocalDiskCleaner/", "打包后的 Windows exe 运行目录。"),
        ("README.md", "开发、运行、测试和打包命令说明。"),
        ("LocalDiskCleaner-friend-test.zip", "可发给他人测试的压缩包。"),
    ], widths=[2800, 6560])
    add_heading(doc, "二、运行环境", 1)
    for item in ["操作系统：Windows 10/11。", "开发语言：Python 3.12。", "桌面界面：Tkinter。", "打包工具：PyInstaller。", "AI 服务：DeepSeek API，作为解释层使用。"]:
        add_bullet(doc, item)
    add_heading(doc, "三、启动方式", 1)
    add_heading(doc, "3.1 直接运行源码", 2)
    add_code(doc, "python tools\\desktop_app.py")
    add_heading(doc, "3.2 双击批处理启动", 2)
    add_code(doc, "LocalDiskCleaner.bat")
    add_heading(doc, "3.3 运行 exe", 2)
    add_code(doc, "dist\\LocalDiskCleaner\\LocalDiskCleaner.exe")
    add_heading(doc, "四、基本使用流程", 1)
    for step in ["打开应用后，在顶部选择要扫描的盘符。", "点击“开始扫描”，等待状态栏显示扫描完成。", "查看分类卡片和左侧项目列表，优先关注“可安全清理”和“可清理但需确认”。", "点击“智能深度分析”，让程序进一步扫描人工判断目录并生成 AI 分析说明。", "阅读右侧分析说明，确认项目含义与风险。", "对绿色项目可点击“移入回收站”；对其他项目建议点击“打开位置”后人工判断。"]:
        add_number(doc, step)
    add_heading(doc, "五、打包方式", 1)
    add_code(doc, "powershell -ExecutionPolicy Bypass -File packaging\\build_exe.ps1")
    add_p(doc, "打包输出位于 dist\\LocalDiskCleaner\\LocalDiskCleaner.exe。若需要发给朋友测试，可压缩 dist\\LocalDiskCleaner 目录或使用已生成的 LocalDiskCleaner-friend-test.zip。")
    add_heading(doc, "六、注意事项", 1)
    for item in ["程序不会永久删除文件，清理动作默认移入回收站。", "黄色、红色和其他项目不会获得回收站清理按钮。", "如果文件被占用，程序会跳过该文件并返回部分成功结果。", "AI 只提供解释建议，不直接决定删除权限。", "公开发布前不建议把 API Key 直接随 exe 分发，当前方案仅适合小范围测试。"]:
        add_bullet(doc, item)
    doc.save(OUT / "02-local-disk-cleaner-run-guide.docx")


def build_test_doc() -> None:
    doc = setup_doc("测试记录与验收文档")
    add_cover(doc, "测试记录与验收文档", "本地磁盘清理器 Local Disk Cleaner")
    add_heading(doc, "一、测试环境", 1)
    add_table(doc, ["项目", "内容"], [
        ("操作系统", "Windows 11"),
        ("Python 环境", "Python 3.12 / Anaconda 开发环境"),
        ("测试框架", "unittest"),
        ("最近一次测试命令", "python -m unittest discover -s tests -v"),
        ("最近一次测试结果", "Ran 63 tests in 1.659s - OK"),
    ], widths=[2600, 6760])
    add_heading(doc, "二、测试用例汇总", 1)
    add_table(doc, ["编号", "测试模块", "关键测试点", "预期结果", "实际结果"], [
        ("T01", "scanner.py", "磁盘根目录规范化、热点目录选择、自定义 target 解析", "返回正确扫描目标", "通过"),
        ("T02", "scanner.py", "递归统计文件和目录大小，跳过过小项目", "统计准确且排序正确", "通过"),
        ("T03", "analyzer.py", "绿色缓存规则匹配", "生成绿色项和回收站 action", "通过"),
        ("T04", "analyzer.py", "用户数据目录分类", "桌面/下载/文档进入人工判断", "通过"),
        ("T05", "cleanup.py", "伪造 action_id、红黄项回收站动作拒绝", "抛出安全错误", "通过"),
        ("T06", "cleanup.py", "缺失路径和占用文件处理", "跳过并返回 partial_success 或 skipped", "通过"),
        ("T07", "ai_deep_scan.py", "深扫黄色父目录并合并候选", "绿色候选从黄色父级扣除空间", "通过"),
        ("T08", "model_explainer.py", "模型 JSON 容错和权限隔离", "文本可解析且不含删除权限", "通过"),
        ("T09", "desktop_app.py", "分类卡片、盘符选择、按钮文案、说明格式", "界面模型输出符合产品术语", "通过"),
        ("T10", "packaging", "exe help 启动和 zip 包检查", "主程序可启动且压缩包含 exe", "通过"),
    ], widths=[900, 1700, 3300, 2100, 1360])
    add_heading(doc, "三、测试命令记录", 1)
    add_code(doc, "python -m unittest discover -s tests -v")
    add_code(doc, "python -m py_compile tools\\scanner.py tools\\analyzer.py tools\\cleanup.py tools\\report.py tools\\pipeline.py tools\\server.py tools\\app.py tools\\webapp.py tools\\desktop_app.py tools\\model_explainer.py tools\\ai_deep_scan.py tools\\secret_store.py tools\\bundle_deepseek_key.py tools\\bundled_secrets.py")
    add_code(doc, "dist\\LocalDiskCleaner\\LocalDiskCleaner.exe --help")
    add_heading(doc, "四、问题与优化记录", 1)
    add_table(doc, ["问题", "原因分析", "解决方案"], [
        ("NVIDIA DXCache 文件被占用导致 WinError 32", "显卡或游戏进程正在使用缓存文件", "清理层跳过占用文件，返回部分成功并建议关闭应用后重试。"),
        ("模型返回内容不是有效 JSON", "LLM 偶发输出不严格符合 JSON 格式", "解析层增加文本回退逻辑，保证 AI 解释失败不影响本地扫描和清理。"),
        ("黄色目录过大但缺少细分建议", "初始扫描只统计父目录", "增加智能深度扫描，把子项重新整合为可安全清理或可清理但需确认。"),
        ("用户难以区分 AI 建议和可执行清理项", "AI 说明位置和主列表关系不清楚", "将 AI 分析作为解释层放在右侧主要阅读区，清理权限仍由本地规则决定。"),
        ("脚本运行方式不适合普通用户", "用户需要命令行和浏览器报告", "改造为 Tkinter 本地桌面应用，并使用 PyInstaller 打包 exe。"),
    ], widths=[2300, 3300, 3760])
    add_heading(doc, "五、验收结论", 1)
    add_p(doc, "从测试结果看，项目核心功能能够正常运行：扫描只读、分类清晰、清理动作受控、AI 只做解释、桌面应用可打包分发。项目满足课程设计中关于需求分析、系统设计、专项开发、智能工具应用、测试验证和成果提交的要求。")
    doc.save(OUT / "03-local-disk-cleaner-test-record.docx")


def build_git_doc() -> None:
    doc = setup_doc("Git版本管理与迭代记录")
    add_cover(doc, "Git版本管理与迭代记录", "本地磁盘清理器 Local Disk Cleaner")
    add_heading(doc, "一、版本管理说明", 1)
    add_p(doc, "任务书要求提交 Git 版本管理日志、迭代记录截图及 git 链接。本项目目录已初始化为 Git 工作区，但当前本地环境未配置远程仓库地址，git log 未返回可展示的提交记录。因此本文档整理实际开发过程中的阶段性迭代内容，提交前可补充 GitHub/Gitee 链接和提交历史截图。")
    add_table(doc, ["检查项", "当前状态"], [
        ("本地 Git 仓库", "存在 .git 目录"),
        ("远程仓库链接", "当前未配置 remote，提交前可补充 GitHub/Gitee 链接"),
        ("git log 输出", "当前环境未返回提交记录"),
        ("建议补交材料", "Git 提交截图、远程仓库链接、主要分支说明"),
    ], widths=[2600, 6760])
    add_heading(doc, "二、阶段性迭代记录", 1)
    add_table(doc, ["阶段", "主要目标", "完成内容"], [
        ("V0 脚本分析原型", "复刻 Storage Analyzer 思路", "实现磁盘扫描、分类分析、HTML 报告生成和本地服务清理动作。"),
        ("V1 规则驱动本地清理软件", "不接大模型，先建立安全边界", "实现绿色/黄色/红色/其他分类、回收站清理、action_id 校验和单元测试。"),
        ("V1.5 桌面化与 exe 打包", "从网页报告改为本地软件", "新增 Tkinter 桌面 UI，支持盘符选择、分类卡片、项目详情和 PyInstaller 打包。"),
        ("V2 AI 解释层", "加入 DeepSeek 但不授予删除权限", "模型负责解释和建议，本地规则负责权限，增加 JSON 容错和路径脱敏。"),
        ("V2.1 智能深度扫描", "让 AI 分析融入扫描结果", "深扫人工判断目录，把可验证缓存并入绿色，把需确认候选并入“可清理但需确认”。"),
        ("V2.2 UI 优化", "提升用户体验", "优化盘符选择、顶部操作区、分类卡片、右侧分析说明和项目详情阅读体验。"),
    ], widths=[1600, 2600, 5160])
    add_heading(doc, "三、建议 Git 提交流程", 1)
    for command in ["git init", "git add .", "git commit -m \"init local disk cleaner\"", "git branch -M main", "git remote add origin <你的GitHub或Gitee仓库地址>", "git push -u origin main"]:
        add_code(doc, command)
    add_heading(doc, "四、工程化规范体现", 1)
    for item in ["模块化代码结构：scanner、analyzer、cleanup、desktop_app、model_explainer 等职责分离。", "自动化测试：tests 目录覆盖扫描、分类、清理安全、AI、桌面 UI、Web 服务等模块。", "可重复打包：packaging/build_exe.ps1 固化 Windows exe 构建流程。", "运行说明完整：README.md 提供开发运行、测试、打包和安全说明。", "安全设计可追溯：需求和设计文档保存在 docs 目录。"]:
        add_bullet(doc, item)
    doc.save(OUT / "04-local-disk-cleaner-git-iteration-record.docx")


if __name__ == "__main__":
    for builder in [build_report, build_run_guide, build_test_doc, build_git_doc]:
        builder()
    aliases = {
        "01-local-disk-cleaner-course-report.docx": "01-本地磁盘清理器-课程设计报告.docx",
        "02-local-disk-cleaner-run-guide.docx": "02-本地磁盘清理器-项目运行说明.docx",
        "03-local-disk-cleaner-test-record.docx": "03-本地磁盘清理器-测试记录与验收文档.docx",
        "04-local-disk-cleaner-git-iteration-record.docx": "04-本地磁盘清理器-Git版本管理与迭代记录.docx",
    }
    for src_name, dst_name in aliases.items():
        src = OUT / src_name
        dst = OUT / dst_name
        dst.write_bytes(src.read_bytes())
    for path in sorted(OUT.glob("*.docx")):
        print(path)
