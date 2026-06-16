#!/usr/bin/env python3
"""Regenerate course deliverable DOCX files for the latest desktop app build."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "course_deliverables"
IMG = OUT / "验证图片"
DATE = "2026年6月16日"

PRIMARY = {
    "report": OUT / "01-本地磁盘清理器-课程设计报告.docx",
    "guide": OUT / "02-本地磁盘清理器-项目运行说明.docx",
    "test": OUT / "03-本地磁盘清理器-测试记录与验收文档.docx",
    "git": OUT / "04-本地磁盘清理器-Git版本管理与迭代记录.docx",
}

ALIASES = {
    PRIMARY["report"]: OUT / "01-local-disk-cleaner-course-report.docx",
    PRIMARY["guide"]: OUT / "02-local-disk-cleaner-run-guide.docx",
    PRIMARY["test"]: OUT / "03-local-disk-cleaner-test-record.docx",
    PRIMARY["git"]: OUT / "04-local-disk-cleaner-git-iteration-record.docx",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(table)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12 if style_name != "Heading 1" else 16)
        style.paragraph_format.space_after = Pt(6)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("互联网应用开发实践")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("667085")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("475467")


def add_meta_table(doc: Document) -> None:
    rows = [
        ("系（院）", "（请填写）"),
        ("专业班级", "网络工程（请填写）"),
        ("姓名", "（请填写）"),
        ("学号", "（请填写）"),
        ("指导教师", "向华"),
        ("项目名称", "本地磁盘清理器 Local Disk Cleaner"),
        ("文档更新时间", DATE),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "内容"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "F2F4F7")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    set_table_widths(table, [1.6, 4.9])


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("1F3A5F")
    p.add_run("\n" + text)
    set_table_widths(table, [6.5])


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_widths(table, widths)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_image(doc: Document, filename: str, caption: str, width: float = 5.9) -> None:
    path = IMG / filename
    if not path.exists():
        doc.add_paragraph(f"截图缺失：{filename}")
        return
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("667085")


def new_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    style_doc(doc)
    add_title(doc, title, subtitle)
    add_meta_table(doc)
    add_callout(
        doc,
        "项目定位",
        "本项目面向 Windows 本地磁盘空间不足场景，提供盘符选择、并发扫描、规则分类、安全清理、AI 智能深度分析和可打包分发的本地桌面应用。AI 只负责解释和补充建议，删除权限始终由本地安全规则决定。",
    )
    return doc


def build_report() -> None:
    doc = new_doc(
        "本地磁盘清理器课程设计报告",
        "Local Disk Cleaner - 安全边界驱动的本地磁盘分析与清理工具",
    )
    doc.add_heading("一、课程设计目的", level=1)
    doc.add_paragraph("本课程设计围绕互联网应用开发实践中的工程化开发、数据处理、安全交互和智能辅助功能展开。项目从磁盘空间不足这一真实问题出发，完成从脚本原型到本地桌面软件的迭代。")
    add_bullets(doc, [
        "掌握本地文件系统扫描、异常处理和数据聚合方法。",
        "通过规则引擎建立清理安全边界，避免误删系统文件和用户重要资料。",
        "使用 Tkinter 和 PyInstaller 完成纯本地桌面应用与 exe 打包。",
        "接入 DeepSeek 作为解释层，实现面向用户的智能分析说明和追问能力。",
    ])

    doc.add_heading("二、设计任务及要求", level=1)
    add_table(doc, ["任务要求", "本项目对应实现"], [
        ["项目选题与需求分析", "围绕 Windows 磁盘空间不足场景，明确普通用户、朋友测试用户和课程验收的使用需求。"],
        ["系统方案创新设计", "设计扫描器、规则分析器、安全清理层、桌面 UI、AI 深度分析解释层和打包分发流程。"],
        ["专项开发任务", "完成数据处理与可视化、轻量化本地工具、智能功能拓展三类任务。"],
        ["工程化规范", "使用 Python 模块化开发、unittest 自动化测试、PyInstaller 打包和 Git 工作区管理。"],
        ["测试与优化", "围绕扫描、分类、权限、防伪 action_id、AI JSON 容错、桌面 UI、并发扫描等建立测试。"],
    ], [2.0, 4.5])

    doc.add_heading("三、需求分析", level=1)
    doc.add_heading("3.1 用户与场景", level=2)
    doc.add_paragraph("目标用户是 Windows 电脑空间不足但不熟悉系统目录的普通用户。用户需要知道哪些文件占空间、哪些可以安全清理、哪些必须人工判断，并希望通过 exe 直接打开软件而不是运行脚本或网页服务。")
    doc.add_heading("3.2 功能需求", level=2)
    add_bullets(doc, [
        "支持选择 C/D/E 等盘符并扫描对应磁盘或系统热点目录。",
        "按照“可安全清理、可清理但需确认、需要人工判断、谨慎处理、其他占用”展示结果。",
        "绿色项目可移入回收站，黄色和红色项目不能直接删除，只能打开位置或查看说明。",
        "智能深度分析可继续扫描人工判断目录，把可验证缓存并入绿色，把候选文件并入“可清理但需确认”。",
        "AI 追问结果显示在右侧“分析说明”区域，帮助用户判断大文件和目录用途。",
    ])
    doc.add_heading("3.3 非功能需求与约束", level=2)
    add_bullets(doc, [
        "扫描过程只读，清理默认移入回收站，不做永久删除。",
        "模型不能授予删除权限，删除权限只来自本地规则和 action_id 校验。",
        "界面必须符合本地软件使用习惯，核心信息区域清晰可读。",
        "扫描速度需要可接受，最新版通过顶层子项并发扫描降低等待时间。",
    ])

    doc.add_heading("四、总体设计", level=1)
    add_table(doc, ["模块", "主要文件", "职责"], [
        ["扫描模块", "tools/scanner.py", "枚举盘符和目标目录，使用 os.scandir 与顶层并发统计大小，记录权限拒绝、占用、跳过重解析点等情况。"],
        ["规则分析模块", "tools/analyzer.py", "按本地规则生成五类项目，并为允许操作的项目生成 action_id。"],
        ["清理安全模块", "tools/cleanup.py", "解析和校验 action_id，只允许本地规则批准的绿色项目移入回收站。"],
        ["桌面应用模块", "tools/desktop_app.py", "提供原生 Tkinter 窗口、盘符选择、分类列表、项目详情、分析说明和操作按钮。"],
        ["AI 分析模块", "tools/ai_deep_scan.py, tools/model_explainer.py", "深扫人工判断目录，调用 DeepSeek 生成解释、优先级建议和项目洞察。"],
        ["打包模块", "packaging/build_exe.ps1", "使用 PyInstaller 生成 dist/LocalDiskCleaner/LocalDiskCleaner.exe 和测试压缩包。"],
    ], [1.1, 1.9, 3.5])
    doc.add_heading("4.1 数据流设计", level=2)
    add_numbered(doc, [
        "用户选择盘符并点击“开始扫描”。",
        "扫描器只读遍历目标目录，按大小输出候选子项。",
        "规则分析器对扫描结果分类并生成可执行动作。",
        "界面展示分类汇总、项目详情和分析说明。",
        "用户点击“智能深度分析”后，系统深扫黄色目录并调用 AI 解释层。",
        "用户根据 AI 建议和本地分类，打开位置或对绿色项执行移入回收站。",
    ])

    doc.add_heading("五、详细设计与实现", level=1)
    doc.add_heading("5.1 只读扫描与速度优化", level=2)
    doc.add_paragraph("最新版扫描器保留精确扫描策略，同时增加顶层子项并发计算。扫描单个大盘符时，多个一级目录可以并行统计，避免一个大型游戏库或 AppData 目录阻塞整个结果生成。扫描接口增加 workers、max_depth 和 progress_callback 参数，为后续快速扫描模式和进度条提供扩展点。")
    doc.add_heading("5.2 规则分类与安全清理", level=2)
    add_table(doc, ["分类", "含义", "允许动作"], [
        ["可安全清理", "已通过本地安全校验的缓存或临时文件", "打开位置、移入回收站"],
        ["可清理但需确认", "AI 深扫发现的候选，但需要用户确认用途", "打开位置"],
        ["需要人工判断", "桌面、文档、下载、聊天软件数据等用户文件", "打开位置"],
        ["谨慎处理", "Windows、Program Files、ProgramData 等系统或程序目录", "查看说明或打开位置"],
        ["其他占用", "已统计但没有明确清理建议", "查看说明"],
    ], [1.35, 3.2, 1.95])
    doc.add_heading("5.3 AI 智能深度分析", level=2)
    doc.add_paragraph("AI 层不直接决定删除权限，而是根据扫描结果、路径特征、大小、分类和用户追问生成解释。智能深度分析会把黄色目录中的子项重新整合到主列表：本地确认的缓存进入绿色，未通过本地安全规则但可能可处理的项目进入“可清理但需确认”。")
    doc.add_heading("5.4 桌面界面实现", level=2)
    doc.add_paragraph("最新版界面将右侧区域调整为上方“项目详情”、下方“大面积分析说明”。用户追问 AI 后，回答追加到分析说明底部并自动滚动，突出 AI 对当前文件的判断依据和手动检查建议。")
    add_image(doc, "应用截图.png", "图 1  本地磁盘清理器主界面与分类结果")

    doc.add_heading("六、系统测试与结果分析", level=1)
    add_table(doc, ["测试类别", "覆盖内容", "最新版结果"], [
        ["扫描测试", "盘符规范化、系统盘热点、数据盘根目录、并发扫描、深度限制、进度回调", "通过"],
        ["规则分类测试", "绿色缓存、用户数据、系统目录、未知路径分类", "通过"],
        ["清理安全测试", "伪造 action_id、红黄项拒绝、缺失路径、占用文件处理", "通过"],
        ["AI 分析测试", "深扫候选合并、模型 JSON 容错、路径脱敏、权限隔离", "通过"],
        ["桌面 UI 测试", "盘符选择、分类卡片、项目详情、分析说明、追问自动滚动", "通过"],
        ["打包验证", "生成 LocalDiskCleaner.exe，更新 friend-test 压缩包", "通过"],
    ], [1.45, 3.9, 1.15])
    doc.add_paragraph("最近一次自动化测试结果：python -m unittest discover -s tests -v，Ran 77 tests，OK。")
    add_image(doc, "清理前D盘.png", "图 2  D 盘清理前空间状态", width=4.8)
    add_image(doc, "清理前E盘.png", "图 3  E 盘清理前空间状态", width=4.8)
    add_image(doc, "清理后D盘和E盘.png", "图 4  D/E 盘清理后空间状态", width=4.8)

    doc.add_heading("七、AI工具与Git工程化实践", level=1)
    doc.add_paragraph("开发过程中参考 Storage Analyzer 技能的思路，将“扫描-分类-报告-安全清理”改造成可交互本地桌面应用。AI 层使用 DeepSeek 作为解释层，帮助用户理解大文件目录，但不会改变本地删除权限。")
    doc.add_paragraph("项目使用 Git 工作区管理开发过程，配合 unittest 进行回归验证。当前交付包包含源码、测试、打包脚本、exe 运行目录和课程文档。")

    doc.add_heading("八、课程设计小结", level=1)
    doc.add_paragraph("本项目完成了从磁盘清理脚本到本地桌面软件的完整迭代。最新版支持盘符选择、并发扫描、智能深度分析、AI 追问、清理权限隔离和 exe 打包，能够让普通用户在较安全的边界内理解和处理磁盘空间占用问题。")
    doc.save(PRIMARY["report"])


def build_guide() -> None:
    doc = new_doc("项目运行说明", "本地磁盘清理器 Local Disk Cleaner")
    doc.add_heading("一、项目文件结构", level=1)
    add_table(doc, ["路径", "说明"], [
        ["tools/", "核心程序模块，包括扫描、分析、清理、AI、桌面 UI 和报告服务。"],
        ["tests/", "单元测试与测试夹具，最新版共 77 项测试。"],
        ["packaging/", "PyInstaller 打包脚本和 spec 文件。"],
        ["dist/LocalDiskCleaner/", "打包后的 Windows exe 运行目录。"],
        ["LocalDiskCleaner-friend-test.zip", "可发给同学或朋友测试的压缩包。"],
        ["course_deliverables/验证图片/", "课程报告使用的软件截图和测试验证截图。"],
    ], [2.1, 4.4])

    doc.add_heading("二、运行环境", level=1)
    add_bullets(doc, [
        "操作系统：Windows 10/11。",
        "开发语言：Python 3.12。",
        "桌面界面：Tkinter。",
        "打包工具：PyInstaller。",
        "AI 服务：DeepSeek API，仅作为解释层和建议层。",
    ])

    doc.add_heading("三、启动方式", level=1)
    doc.add_heading("3.1 直接运行源码", level=2)
    doc.add_paragraph(r"python tools\desktop_app.py")
    doc.add_heading("3.2 指定扫描参数", level=2)
    doc.add_paragraph(r"python tools\desktop_app.py --workers 8 --limit 80")
    doc.add_paragraph("workers 用于控制扫描并发度；普通用户不需要修改，默认会自动选择合理值。")
    doc.add_heading("3.3 运行 exe", level=2)
    doc.add_paragraph(r"dist\LocalDiskCleaner\LocalDiskCleaner.exe")
    doc.add_heading("3.4 分发测试包", level=2)
    doc.add_paragraph(r"将 LocalDiskCleaner-friend-test.zip 解压后，双击 LocalDiskCleaner.exe 即可启动。")

    doc.add_heading("四、基本使用流程", level=1)
    add_numbered(doc, [
        "打开软件后，在顶部下拉框选择需要扫描的盘符。",
        "点击“开始扫描”，等待左侧分类列表和顶部汇总卡片更新。",
        "优先查看“可安全清理”和“可清理但需确认”分类。",
        "选中项目后，在右侧上方查看项目详情，可打开位置或对绿色项移入回收站。",
        "点击“智能深度分析”，系统会继续扫描人工判断目录，并在右侧下方“分析说明”生成 AI 建议。",
        "如仍不确定，可在追问框输入问题，AI 回答会追加到分析说明区域底部。",
    ])

    doc.add_heading("五、打包方式", level=1)
    doc.add_paragraph(r"powershell -ExecutionPolicy Bypass -File packaging\build_exe.ps1")
    doc.add_paragraph("打包完成后会生成 dist/LocalDiskCleaner/LocalDiskCleaner.exe。若 dist 目录被旧进程锁定，需要先关闭正在运行的 LocalDiskCleaner.exe。")

    doc.add_heading("六、注意事项", level=1)
    add_bullets(doc, [
        "清理操作默认移入回收站，清空回收站前仍可恢复。",
        "AI 只解释和建议，不会授予删除权限。",
        "黄色、红色和“可清理但需确认”项目需要人工确认，软件不会自动删除。",
        "若遇到文件占用，可关闭相关软件后重新尝试。",
        "DeepSeek API Key 已按测试版要求做应用内简单加密，适合课程和朋友测试，不应作为正式商业密钥保护方案。",
    ])
    doc.save(PRIMARY["guide"])


def build_test_record() -> None:
    doc = new_doc("测试记录与验收文档", "本地磁盘清理器 Local Disk Cleaner")
    doc.add_heading("一、测试环境", level=1)
    add_table(doc, ["项目", "内容"], [
        ["操作系统", "Windows 11"],
        ["Python 环境", "Python 3.12 / Anaconda 开发环境"],
        ["测试框架", "unittest"],
        ["最近一次测试命令", "python -m unittest discover -s tests -v"],
        ["最近一次测试结果", "Ran 77 tests in 1.655s - OK"],
        ["打包验证", "LocalDiskCleaner.exe --help 可正常输出参数，压缩包包含主程序。"],
    ], [1.8, 4.7])

    doc.add_heading("二、测试用例汇总", level=1)
    add_table(doc, ["编号", "测试模块", "关键测试点", "实际结果"], [
        ["T01", "scanner.py", "盘符规范化、系统盘热点目录、数据盘根目录、自定义 target 解析", "通过"],
        ["T02", "scanner.py", "递归统计、并发扫描 workers、max_depth 深度限制、progress_callback 进度事件", "通过"],
        ["T03", "analyzer.py", "绿色缓存、黄色用户数据、红色系统路径、未知路径分类", "通过"],
        ["T04", "cleanup.py", "伪造 action_id 拒绝、红黄项回收站动作拒绝、缺失路径和占用文件处理", "通过"],
        ["T05", "ai_deep_scan.py", "深扫黄色父目录并把候选合并为绿色或可清理但需确认", "通过"],
        ["T06", "model_explainer.py", "模型 JSON 容错、路径脱敏、AI 解释不改变本地权限", "通过"],
        ["T07", "desktop_app.py", "盘符选择、项目详情、分析说明放大、追问结果自动滚动到底部", "通过"],
        ["T08", "packaging", "PyInstaller exe 构建与 friend-test 压缩包验证", "通过"],
    ], [0.55, 1.25, 3.9, 0.8])

    doc.add_heading("三、测试命令记录", level=1)
    doc.add_paragraph("python -m unittest discover -s tests -v")
    doc.add_paragraph(r"dist\LocalDiskCleaner\LocalDiskCleaner.exe --help")
    doc.add_paragraph(r"Compress-Archive -Path .\dist\LocalDiskCleaner\* -DestinationPath .\LocalDiskCleaner-friend-test.zip -Force")

    doc.add_heading("四、问题与优化记录", level=1)
    add_table(doc, ["问题", "原因分析", "解决方案"], [
        ["NVIDIA DXCache 文件被占用导致 WinError 32", "显卡或游戏进程正在使用缓存文件", "清理层跳过占用文件，返回部分成功并建议关闭应用后重试。"],
        ["模型返回内容不是有效 JSON", "LLM 偶发输出不严格符合 JSON 格式", "解析层增加文本回退逻辑，保证 AI 解释失败不影响本地扫描和清理。"],
        ["黄色目录过大但缺少细分建议", "初始扫描只统计父目录", "增加智能深度扫描，把子项重新整合为绿色或可清理但需确认。"],
        ["扫描速度偏慢", "旧版逐个目录递归统计，一个大目录会阻塞后续结果", "扫描器增加顶层并发 workers、可选 max_depth 和进度回调扩展点。"],
        ["AI 追问结果不够显眼", "分析说明和项目详情区域布局不符合阅读重点", "将项目详情上移并压缩，把分析说明下移并放大。"],
        ["脚本运行方式不适合普通用户", "用户需要命令行和浏览器报告", "改造为 Tkinter 本地桌面应用，并使用 PyInstaller 打包 exe。"],
    ], [1.65, 2.15, 2.7])

    doc.add_heading("五、截图验证", level=1)
    add_image(doc, "应用截图.png", "图 1  软件主界面截图")
    add_image(doc, "报告清理后截图.png", "图 2  清理后报告/结果截图")
    add_image(doc, "同学1电脑测试可清除数据.png", "图 3  同学 1 电脑测试可清理数据截图")
    add_image(doc, "同学2电脑测试可清除数据.png", "图 4  同学 2 电脑测试可清理数据截图")
    add_image(doc, "清理后D盘和E盘.png", "图 5  D/E 盘清理后空间状态", width=4.8)

    doc.add_heading("六、验收结论", level=1)
    doc.add_paragraph("本项目已完成最新版功能验收：可选择盘符扫描，支持并发扫描优化，能够按安全等级展示空间占用，绿色项目可移入回收站，AI 智能深度分析和追问仅提供解释建议，不改变本地删除权限。自动化测试 77 项通过，exe 和测试压缩包已重新生成。")
    doc.save(PRIMARY["test"])


def build_git_record() -> None:
    doc = new_doc("Git版本管理与迭代记录", "本地磁盘清理器 Local Disk Cleaner")
    doc.add_heading("一、版本管理说明", level=1)
    doc.add_paragraph("本项目已上传到 GitHub 公开仓库，仓库用于保存项目源码、README、打包目录和阶段性提交记录。本文档根据远程仓库信息和本地开发记录整理版本管理过程。")
    add_table(doc, ["检查项", "当前状态"], [
        ["本地 Git 仓库", "存在 .git 目录"],
        ["远程仓库链接", "https://github.com/ciccicicii/CCleanerX"],
        ["仓库可见性", "Public"],
        ["默认分支", "main"],
        ["仓库创建时间", "2026-06-16 12:34:31（北京时间，GitHub API 返回 UTC 04:34:31）"],
        ["最近推送时间", "2026-06-16 17:40:04（北京时间，GitHub API 返回 UTC 09:40:04）"],
        ["仓库主要目录", "LocalDiskCleaner/、README.md"],
        ["自动化测试", "最近一次运行 77 项 unittest，全部通过"],
        ["可分发产物", "dist/LocalDiskCleaner/LocalDiskCleaner.exe 与 LocalDiskCleaner-friend-test.zip"],
        ["验证图片", r"course_deliverables\验证图片"],
    ], [1.9, 4.6])

    doc.add_heading("二、远程提交记录", level=1)
    add_table(doc, ["提交号", "提交时间（北京时间）", "提交说明", "对应迭代意义"], [
        ["657180c", "2026-06-16 12:35", "first commit", "建立远程仓库和项目初始版本。"],
        ["a85f893", "2026-06-16 12:37", "初步测试", "完成初始功能验证，形成可测试原型。"],
        ["7f4b041", "2026-06-16 12:53", "可选用多磁盘进行分析", "从固定 C 盘扩展为可选择不同盘符扫描。"],
        ["7fd534a", "2026-06-16 13:09", "分出智能分析框调整为可选中项目", "开始强化 AI 分析区域与项目选择之间的交互。"],
        ["23e4f89", "2026-06-16 13:28", "添加分析说明框", "增加独立分析说明区域，承载 AI 解释和清理建议。"],
        ["bae24a5", "2026-06-16 13:35", "界面优化", "优化桌面界面布局、分类卡片和用户阅读体验。"],
        ["e50d2df", "2026-06-16 16:14", "对AI解释层进行增强...", "为 AI 投喂盘符、扩展名、风险、置信度等多维特征，并加入 AI 追问功能。"],
        ["df56fb3", "2026-06-16 17:10", "修复AI返回数据处理带隐藏字段", "修复 AI 返回结果中隐藏字段对用户可读性的影响。"],
        ["ce6795a", "2026-06-16 17:39", "优化扫描性能", "增加扫描性能优化，改善大盘符和大目录场景下的等待时间。"],
    ], [0.8, 1.35, 2.2, 2.15])

    doc.add_heading("三、阶段性迭代归纳", level=1)
    add_table(doc, ["阶段", "主要目标", "完成内容"], [
        ["V0 脚本分析原型", "复刻 Storage Analyzer 思路", "实现磁盘扫描、分类分析、HTML 报告生成和本地服务清理动作。"],
        ["V1 规则驱动本地清理软件", "不接大模型，先建立安全边界", "实现绿色/黄色/红色/其他分类、回收站清理、action_id 校验和单元测试。"],
        ["V1.5 桌面化与 exe 打包", "从网页报告改为本地软件", "新增 Tkinter 桌面 UI，支持盘符选择、分类卡片、项目详情和 PyInstaller 打包。"],
        ["V2 AI 解释层", "加入 DeepSeek 但不授予删除权限", "模型负责解释和建议，本地规则负责权限，增加 JSON 容错和路径脱敏。"],
        ["V2.1 智能深度扫描", "让 AI 分析融入扫描结果", "深扫人工判断目录，把可验证缓存并入绿色，把需确认候选并入“可清理但需确认”。"],
        ["V2.2 UI 优化", "提升用户体验", "优化盘符选择、顶部操作区、分类卡片、右侧分析说明和项目详情阅读体验。"],
        ["V2.3 朋友测试版", "便于他人测试", "将 API Key 做应用内简单加密，重新打包 exe 和测试压缩包。"],
        ["V2.4 扫描性能优化", "减少大盘符扫描等待", "顶层子项并发扫描，增加 workers、max_depth、progress_callback，并将普通扫描和智能深扫接入默认并发。"],
        ["V2.5 AI 阅读区优化", "让追问结果更容易看见", "将项目详情上移并压缩，把分析说明移到下方并扩大，追问结果追加后自动滚动到底部。"],
    ], [1.15, 1.75, 3.6])

    doc.add_heading("四、Git 操作流程", level=1)
    doc.add_paragraph("git init")
    doc.add_paragraph("git add .")
    doc.add_paragraph('git commit -m \"feat: complete local disk cleaner desktop app\"')
    doc.add_paragraph("git remote add origin https://github.com/ciccicicii/CCleanerX.git")
    doc.add_paragraph("git push -u origin main")

    doc.add_heading("五、工程化规范体现", level=1)
    add_bullets(doc, [
        "模块边界清晰：扫描、分析、清理、AI、UI、报告和打包分离。",
        "安全约束清晰：AI 不能授予删除权限，清理动作必须由本地规则生成。",
        "测试覆盖关键路径：扫描、分类、清理权限、AI 容错、桌面 UI 和打包入口均有验证。",
        "远程提交记录完整：GitHub 仓库显示 2026-06-16 形成 9 次阶段性提交，能对应多磁盘扫描、AI 增强、界面优化和性能优化等关键迭代。",
    ])
    doc.save(PRIMARY["git"])


def copy_aliases() -> None:
    for source, target in ALIASES.items():
        shutil.copyfile(source, target)


def main() -> int:
    OUT.mkdir(exist_ok=True)
    build_report()
    build_guide()
    build_test_record()
    build_git_record()
    copy_aliases()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
